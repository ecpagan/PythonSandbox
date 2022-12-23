# pip install python-docx is required for this to work
import docx


def get_text(filename):
    doc = docx.Document(filename)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return '\n'.join(content)


docx_path = r''

d = docx.Document(docx_path)
d.add_paragraph('Hello this is a paragraph.')
d.save(docx_path)

print(get_text(docx_path))
